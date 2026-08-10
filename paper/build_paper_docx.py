"""Build the IEI full-length paper in MS Word format from paper/paper_source.md.

The IEI submission instructions make the .doc/.docx file mandatory and the PDF
a companion copy, so this script and build_paper.py read the SAME source file.
paper_source.md is the single point of truth; neither output is hand-edited.

Formatting follows the conventional IEI author guidelines: A4, Times New Roman,
11 pt body, single column, justified body text, numbered headings, figures and
tables captioned in sequence.

Usage:  python paper/build_paper_docx.py
Output: paper/AgriRover_IEI_38th_National_Convention_Full_Paper.docx
"""

import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from build_paper import TABLE_TITLES, parse  # single source of truth for both

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "paper_source.md")
FIGDIR = os.path.join(HERE, "figures")
OUT = os.path.join(
    HERE, "AgriRover_IEI_38th_National_Convention_Full_Paper.docx")

BODY_FONT = "Times New Roman"
INK = RGBColor(0x1A, 0x1C, 0x1A)
ACCENT = RGBColor(0x2F, 0x5D, 0x3A)
MUTED = RGBColor(0x5D, 0x66, 0x60)

# Content width inside the 2.54 cm side margins the IEI guidelines specify.
MARGIN_CM = 2.54
CONTENT_W_CM = 21.0 - 2 * MARGIN_CM

# IEI author guidelines: 12 pt body, 1.5 line spacing, 14 pt bold centred
# title, 12 pt bold main headings, 12 pt bold italic sub headings.
BODY_PT = 12
LEADING = 1.5

FIG_WIDTHS = {
    "fig1_architecture.png": 1.00,
    "fig2_safety_chain.png": 1.00,
    "fig3_coverage_path.png": 1.00,
    "fig4_localisation_error.png": 0.74,
}


# ---------------------------------------------------------------- docx utils
def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def cell_margins(table, twips=80):
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tblPr.append(mar)


def para(doc, text_runs, *, size=BODY_PT, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         bold=False, italic=False, color=INK, before=0, after=6,
         leading=LEADING, font=BODY_FONT, keep_with_next=False):
    """text_runs is a plain string or a list of (text, bold, italic) tuples."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = leading
    pf.keep_with_next = keep_with_next
    if isinstance(text_runs, str):
        text_runs = [(text_runs, bold, italic)]
    for txt, b, i, sub, sup in (norm(t) for t in text_runs):
        r = p.add_run(txt)
        r.font.name = font
        r.font.size = Pt(size)
        r.font.bold = b
        r.font.italic = i
        r.font.color.rgb = color
        r.font.subscript = sub or None
        r.font.superscript = sup or None
        # Ensure the East-Asian font slot matches, or Word may substitute.
        r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return p


def hrule(paragraph, color="C8CEC9", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


# ------------------------------------------------------------ inline markup
INLINE_RE = re.compile(r"\*\*(.+?)\*\*")


def norm(t):
    """Pad a run tuple to (text, bold, italic, sub, sup)."""
    return tuple(t) + (False,) * (5 - len(t))


def bold_runs(text):
    text = text.replace("--", "\u2014")
    out, pos = [], 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], False, False, False, False))
        out.append((m.group(1), True, False, False, False))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False, False, False, False))
    return out


def inline_math_runs(tex):
    """Runs for a short in-prose expression such as $F_k$: italic variables
    with real subscripts, rather than a literal underscore on the page."""
    out = []
    for k, v in MATH_SYMBOLS.items():
        tex = tex.replace(k, v)
    pos = 0
    for m in re.finditer(r"\^\{([^{}]*)\}|_\{([^{}]*)\}|\^(.)|_(.)", tex):
        if m.start() > pos:
            out.append((tex[pos:m.start()], False, True, False, False))
        sup = m.group(1) if m.group(1) is not None else m.group(3)
        sub = m.group(2) if m.group(2) is not None else m.group(4)
        if sup is not None:
            out.append((sup, False, True, False, True))
        else:
            out.append((sub, False, True, True, False))
        pos = m.end()
    if pos < len(tex):
        out.append((tex[pos:], False, True, False, False))
    return out


def runs_from(text):
    """Split **bold** and $math$ spans into run tuples."""
    out = []
    for i, part in enumerate(re.split(r"\$(.+?)\$", text)):
        out.extend(inline_math_runs(part) if i % 2 else bold_runs(part))
    return out or [("", False, False, False, False)]


# ------------------------------------------------------------------- math
MATH_SYMBOLS = {
    r"\Delta": "\u0394", r"\theta": "\u03b8", r"\omega": "\u03c9",
    r"\top": "\u1d40", r"\times": "\u00d7", r"\cdot": "\u00b7",
    r"\sigma": "\u03c3", r"\mu": "\u03bc", r"\phi": "\u03c6",
}
SUP = {"0": "\u2070", "1": "\u00b9", "2": "\u00b2", "3": "\u00b3",
       "-": "\u207b", "T": "\u1d40"}
SUB_DIGITS = {"0": "\u2080", "1": "\u2081", "2": "\u2082", "3": "\u2083",
              "4": "\u2084", "5": "\u2085", "6": "\u2086", "7": "\u2087",
              "8": "\u2088", "9": "\u2089", "k": "\u2096", "|": "|",
              "-": "\u208b"}


def math_paragraph(doc, tex):
    """Render display math as a centred paragraph with real sub/superscripts.

    Word has no plain-text way to italicise variables while keeping operator
    names upright, so we build the equation run by run instead of as a string.
    """
    t = tex.strip()
    t = t.replace(r"\qquad", "\u2003\u2003").replace(r"\quad", "\u2003")
    t = t.replace(r"\left(", "(").replace(r"\right)", ")")
    t = t.replace(r"\left[", "[").replace(r"\right]", "]")
    t = t.replace("\\,", "\u2009").replace("\\;", "\u200a")
    for k, v in MATH_SYMBOLS.items():
        t = t.replace(k, v)
    for fn in ("cos", "sin", "tan", "exp", "log"):
        t = t.replace("\\" + fn, f"\x00{fn}\x00\u2009")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)

    def emit(txt, *, sub=False, sup=False, italic=True):
        if not txt:
            return
        r = p.add_run(txt)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.italic = italic
        r.font.color.rgb = INK
        r._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        if sub:
            r.font.subscript = True
        if sup:
            r.font.superscript = True

    # \x00-delimited spans are upright operator names.
    token = re.compile(r"\x00(?P<fn>[a-z]+)\x00"
                       r"|\^\{(?P<supb>[^{}]*)\}"
                       r"|_\{(?P<subb>[^{}]*)\}"
                       r"|\^(?P<sup1>.)"
                       r"|_(?P<sub1>.)")
    pos = 0
    for m in token.finditer(t):
        if m.start() > pos:
            emit(t[pos:m.start()])
        g = m.groupdict()
        if g["fn"]:
            emit(g["fn"], italic=False)
        elif g["supb"] is not None:
            emit(g["supb"], sup=True)
        elif g["subb"] is not None:
            emit(g["subb"], sub=True)
        elif g["sup1"]:
            emit(g["sup1"], sup=True)
        elif g["sub1"]:
            emit(g["sub1"], sub=True)
        pos = m.end()
    if pos < len(t):
        emit(t[pos:])
    return p


# ------------------------------------------------------------------ blocks
def figure_block(doc, fname, caption, num):
    path = os.path.join(FIGDIR, fname)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(
        path, width=Cm(CONTENT_W_CM * FIG_WIDTHS.get(fname, 0.92)))
    cap = [(f"Figure {num}. ", True, False)] + runs_from(caption)
    para(doc, cap, size=9, color=MUTED, after=10)


def table_block(doc, rows, num):
    header, body = rows[0], rows[1:]
    ncol = len(header)

    para(doc, TABLE_TITLES.get(num, f"Table {num}."),
         size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
         before=8, after=3, keep_with_next=True)

    table = doc.add_table(rows=1, cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell_margins(table)

    if ncol == 3:
        fracs = [0.30, 0.42, 0.28]
    elif ncol == 4:
        fracs = [0.31, 0.23, 0.23, 0.23]
    else:
        fracs = [1 / ncol] * ncol
    widths = [Cm(CONTENT_W_CM * f) for f in fracs]

    def fill(cells, values, *, head=False, band=False):
        for cell, val, w in zip(cells, values, widths):
            cell.width = w
            if head:
                shade(cell, "2F5D3A")
            elif band:
                shade(cell, "EEF2EF")
            cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for txt, b, i, sub, sup in (norm(t) for t in runs_from(val)):
                r = p.add_run(txt)
                r.font.name = BODY_FONT
                r.font.size = Pt(9)
                r.font.bold = head or b
                r.font.italic = i
                r.font.subscript = sub or None
                r.font.superscript = sup or None
                r.font.color.rgb = (
                    RGBColor(0xFF, 0xFF, 0xFF) if head else INK)
                r._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    fill(table.rows[0].cells, header, head=True)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    for idx, r in enumerate(body, start=1):
        r = (r + [""] * ncol)[:ncol]
        fill(table.add_row().cells, r, band=(idx % 2 == 0))

    para(doc, "", after=8)


# ---------------------------------------------------------------- document
def page_setup(doc):
    sec = doc.sections[0]
    sec.start_type = WD_SECTION.NEW_PAGE
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(MARGIN_CM)
    sec.right_margin = Cm(MARGIN_CM)
    sec.top_margin = Cm(MARGIN_CM)
    sec.bottom_margin = Cm(MARGIN_CM)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.line_spacing = LEADING
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    # Running head, suppressed on the title page by Word's own first-page flag.
    sec.different_first_page_header_footer = True
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("AgriRover  |  38th National Convention of Agricultural "
                    "Engineers, The Institution of Engineers (India)")
    hr.font.name = BODY_FONT
    hr.font.size = Pt(8)
    hr.font.color.rgb = MUTED

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


def add_page_number(paragraph):
    run = paragraph.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._r.append(el)


def masthead(doc, meta):
    # Title is 14 pt bold centred exactly as the guidelines require.
    para(doc, meta["TITLE"], size=14, bold=True, color=INK,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=4, leading=1.15)
    para(doc, meta["SUBTITLE"], size=BODY_PT, italic=True, color=INK,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10, leading=1.15)

    authors = " \u00b7 ".join(a.strip() for a in meta["AUTHORS"].split("|"))
    para(doc, authors, size=BODY_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=3, leading=1.15)
    para(doc, meta["AFFILIATION"], size=10, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2, leading=1.15)
    para(doc, meta["CONTACT"], size=10, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2, leading=1.15)
    para(doc, meta["VENUE"], size=10, italic=True, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10, leading=1.15)

    # Abstract in a single shaded cell, mirroring the PDF layout.
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    cell_margins(t, 120)
    cell = t.rows[0].cells[0]
    cell.width = Cm(CONTENT_W_CM)
    shade(cell, "EEF2EF")
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)

    h = cell.add_paragraph()
    hr = h.add_run("ABSTRACT")
    hr.font.name = BODY_FONT
    hr.font.size = Pt(BODY_PT)
    hr.font.bold = True
    hr.font.color.rgb = ACCENT

    b = cell.add_paragraph()
    b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    b.paragraph_format.line_spacing = LEADING
    for txt, bo, it, sub, sup in (norm(t) for t in runs_from(meta["ABSTRACT"])):
        r = b.add_run(txt)
        r.font.name = BODY_FONT
        r.font.size = Pt(BODY_PT)
        r.font.bold = bo
        r.font.italic = it
        r.font.subscript = sub or None
        r.font.superscript = sup or None
        r.font.color.rgb = INK

    para(doc, [("Keywords: ", True, False)] + runs_from(meta["KEYWORDS"]),
         size=BODY_PT, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=4)


def build():
    meta, blocks = parse(SRC)
    doc = Document()
    page_setup(doc)

    props = doc.core_properties
    props.title = meta["TITLE"]
    props.author = "V. K. Gupta, K. H. Mukeshbhai, S. Wagh, P. Nandy"
    props.subject = ("38th National Convention of Agricultural Engineers, "
                     "The Institution of Engineers (India)")
    props.keywords = meta["KEYWORDS"]

    masthead(doc, meta)

    fig_no = tbl_no = ref_no = 0
    for kind, payload in blocks:
        if kind == "h1":
            # Main heading: 12 pt bold, per the IEI paper format table.
            p = para(doc, payload, size=BODY_PT, bold=True, color=INK,
                     align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=5,
                     leading=1.15, keep_with_next=True)
            hrule(p)
        elif kind == "h2":
            # Sub heading: 12 pt bold italic.
            para(doc, payload, size=BODY_PT, bold=True, italic=True, color=INK,
                 align=WD_ALIGN_PARAGRAPH.LEFT, before=9, after=3,
                 leading=1.15, keep_with_next=True)
        elif kind == "p":
            para(doc, runs_from(payload))
        elif kind == "math":
            math_paragraph(doc, payload)
        elif kind == "fig":
            fig_no += 1
            figure_block(doc, payload[0], payload[1], fig_no)
        elif kind == "table":
            tbl_no += 1
            table_block(doc, payload, tbl_no)
        elif kind == "ref":
            ref_no += 1
            p = para(doc, [(f"[{ref_no}]  ", False, False)]
                     + runs_from(payload), size=9.5, after=3)
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.75)

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
